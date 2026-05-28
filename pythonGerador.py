import redis
import io
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

def format_output(res, indent=0):
    """Formata as respostas do Redis imitando o redis-cli, suportando listas aninhadas (como no SCAN)"""
    prefix = " " * indent
    if res is True:
        return f"{prefix}(integer) 1"
    elif res is False:
        return f"{prefix}(integer) 0"
    elif res is None:
        return f"{prefix}(nil)"
    elif isinstance(res, int):
        return f"{prefix}(integer) {res}"
    elif isinstance(res, bytes):
        return f"{prefix}\"{res.decode('utf-8')}\""
    elif isinstance(res, str):
        if res == 'OK':
            return f"{prefix}{res}"
        return f"{prefix}\"{res}\""
    elif isinstance(res, list):
        if not res:
            return f"{prefix}(empty array)"
        lines = []
        for i, item in enumerate(res):
            if isinstance(item, list):
                # Se for uma lista dentro de lista (Ex: retorno do SCAN)
                lines.append(f"{prefix}{i+1})")
                lines.append(format_output(item, indent + 3))
            else:
                # Formata o item e remove o prefixo extra caso a recursão adicione
                formatted_item = format_output(item).strip() 
                lines.append(f"{prefix}{i+1}) {formatted_item}")
        return "\n".join(lines)
    else:
        return f"{prefix}{str(res)}"

def create_terminal_image(text):
    """Fabrica uma imagem PNG em memória imitando um terminal escuro"""
    try:
        font = ImageFont.load_default(size=16)
        char_width, line_height = 10, 22
    except TypeError:
        font = ImageFont.load_default()
        char_width, line_height = 8, 15

    linhas = text.split('\n')
    max_caracteres = max([len(linha) for linha in linhas] + [1])
    
    largura = (max_caracteres * char_width) + 40
    altura = (len(linhas) * line_height) + 40

    img = Image.new('RGB', (largura, altura), color=(12, 12, 12))
    d = ImageDraw.Draw(img)

    d.multiline_text((20, 20), text, font=font, fill=(204, 204, 204), spacing=4)

    img_stream = io.BytesIO()
    img.save(img_stream, format='PNG')
    img_stream.seek(0)
    
    return img_stream

def main():
    try:
        r = redis.Redis(host='localhost', port=6379, decode_responses=True)
        r.ping()
        print("Conectado ao Redis com sucesso! Gerando documento...")
    except redis.ConnectionError:
        print("Erro: Não foi possível conectar ao Redis na porta 6379.")
        return

    # Estrutura aprimorada:
    # 'setup': Comandos rodados de forma oculta para preparar o banco de dados.
    # 'run': Comandos que de fato vão aparecer no print do Word.
    exercicios = {
        1: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "usuario", "Carlos"], ["GET", "usuario"], ["SET", "usuario", "Carlos Silva"], ["GET", "usuario"]]
        },
        2: {
            "setup": [["FLUSHDB"]], 
            "run": [["MSET", "produto", "Notebook", "preco", "3500", "estoque", "15"], ["MGET", "produto", "preco", "estoque"]]
        },
        3: {
            "setup": [["FLUSHDB"]], 
            "run": [["SETNX", "admin", "true"], ["SETNX", "admin", "false"]]
        },
        4: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "nome", "Maria"], ["APPEND", "nome", " Oliveira"], ["STRLEN", "nome"], ["GETRANGE", "nome", "0", "4"]]
        },
        5: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "cidade", "Campinas"], ["SETRANGE", "cidade", "0", "São "], ["GET", "cidade"]]
        },
        6: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "pontos", "10"], ["INCR", "pontos"], ["INCRBY", "pontos", "5"], ["DECRBY", "pontos", "3"], ["GET", "pontos"]]
        },
        7: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "saldo", "100.50"], ["INCRBYFLOAT", "saldo", "25.75"], ["GET", "saldo"]]
        },
        8: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "token", "xyz123", "EX", "60"], ["TTL", "token"], ["PERSIST", "token"], ["TTL", "token"]]
        },
        9: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "sessao", "ativa"], ["PEXPIRE", "sessao", "5000"], ["PTTL", "sessao"]]
        },
        10: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "curso", "Redis"], ["RENAME", "curso", "disciplina"], ["COPY", "disciplina", "backup_disciplina"], ["TYPE", "disciplina"], ["EXISTS", "disciplina"], ["DEL", "disciplina"]]
        },
        11: {
            "setup": [["FLUSHDB"]], 
            "run": [["SET", "minha_chave", "teste"], ["MOVE", "minha_chave", "1"], ["SELECT", "1"], ["GET", "minha_chave"], ["SELECT", "0"]]
        },
        12: {
            "setup": [["SELECT", "0"], ["FLUSHDB"]], # Limpa tudo antes para mostrar APENAS 5 chaves
            "run": [["MSET", "k1", "1", "k2", "2", "k3", "3", "k4", "4", "k5", "5"], ["KEYS", "*"], ["SCAN", "0"]]
        },
        13: {
            "setup": [["FLUSHDB"]], 
            "run": [["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios", "Dormir"], ["LRANGE", "tarefas", "0", "-1"]]
        },
        14: {
            "setup": [["FLUSHDB"], ["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios", "Dormir"]], 
            "run": [["LLEN", "tarefas"], ["LINDEX", "tarefas", "0"], ["LINDEX", "tarefas", "-1"]]
        },
        15: {
            "setup": [["FLUSHDB"], ["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios", "Dormir"]], 
            "run": [["LSET", "tarefas", "1", "Fazer Exercicios Praticos"], ["LRANGE", "tarefas", "0", "-1"]]
        },
        16: {
            "setup": [["FLUSHDB"], ["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios Praticos", "Dormir"]], 
            "run": [["LINSERT", "tarefas", "BEFORE", "Dormir", "Tomar Cafe"], ["LRANGE", "tarefas", "0", "-1"]]
        },
        17: {
            "setup": [["FLUSHDB"], ["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios Praticos", "Tomar Cafe", "Dormir"]], 
            "run": [["LPOP", "tarefas"], ["RPOP", "tarefas"], ["LRANGE", "tarefas", "0", "-1"]]
        },
        18: {
            "setup": [["FLUSHDB"]], 
            "run": [["RPUSH", "fila_atividade_pendente", "Tarefa 1", "Tarefa 2"], ["LMOVE", "fila_atividade_pendente", "fila_atividade_processando", "LEFT", "RIGHT"]]
        },
        19: {
            "setup": [["FLUSHDB"]], 
            "run": [["RPUSH", "logs", "L1", "L2", "L3", "L4", "L5", "L6", "L7", "L8", "L9", "L10"], ["LTRIM", "logs", "-5", "-1"], ["LRANGE", "logs", "0", "-1"]]
        },
        20: {
            "setup": [["FLUSHDB"], ["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios"]], 
            "run": [["RPUSH", "tarefas", "Dormir", "Dormir", "Dormir"], ["LREM", "tarefas", "1", "Dormir"], ["LREM", "tarefas", "0", "Dormir"], ["LRANGE", "tarefas", "0", "-1"]]
        },
        21: {
            "setup": [["FLUSHDB"], ["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios", "Dormir", "Tomar Cafe", "Dormir"]], 
            "run": [["LPOS", "tarefas", "Estudar Redis"], ["LPOS", "tarefas", "Dormir", "COUNT", "0"]]
        },
        22: {
            "setup": [["FLUSHDB"], ["RPUSH", "tarefas", "Estudar Redis", "Fazer Exercicios", "Estudar Redis", "Dormir"]], 
            "run": [["LREM", "tarefas", "0", "Estudar Redis"], ["DEL", "tarefas"]]
        },
        23: {
            "setup": [["FLUSHDB"]], 
            "run": [["RPUSH", "lista_temp", "item1", "item2"], ["EXPIRE", "lista_temp", "10"], ["TTL", "lista_temp"], ["EXISTS", "lista_temp"]]
        },
        24: {
            "setup": [["FLUSHDB"]], 
            "run": [
                ["SET", "ip:192.168.1.50:tentativas", "1", "EX", "10", "NX"],
                ["INCR", "ip:192.168.1.50:tentativas"],
                ["INCR", "ip:192.168.1.50:tentativas"],
                ["GET", "ip:192.168.1.50:tentativas"],
                ["INCRBY", "ip:192.168.1.50:tentativas", "3"],
                ["SET", "ip:192.168.1.50:bloqueado", "true", "EX", "3600"],
                ["EXISTS", "ip:192.168.1.50:bloqueado"]
            ]
        }
    }

    doc = Document()
    doc.add_heading('Resolução: Exercícios Práticos - Redis CLI', 0)

    for ex_num, block in exercicios.items():
        doc.add_heading(f'Exercício {ex_num}', level=1)
        
        # 1. Executa os comandos de setup de forma silenciosa para isolar o ambiente
        for cmd in block.get("setup", []):
            try:
                r.execute_command(*cmd)
            except Exception:
                pass # Ignora erros no setup

        # 2. Executa os comandos reais que farão parte do print
        output_lines = []
        for cmd in block.get("run", []):
            cmd_str = " ".join(cmd)
            output_lines.append(f"127.0.0.1:6379> {cmd_str}")
            
            try:
                res = r.execute_command(*cmd)
                res_str = format_output(res)
                output_lines.append(res_str)
            except Exception as e:
                output_lines.append(f"(error) {str(e)}")

        texto_terminal = "\n".join(output_lines)
        img_stream = create_terminal_image(texto_terminal)
        
        doc.add_picture(img_stream, width=Inches(5.5))
        print(f"Exercício {ex_num} processado com sucesso.")

    nome_arquivo = "Lista_Exercicios_Redis.docx"
    doc.save(nome_arquivo)
    print(f"\nDocumento '{nome_arquivo}' gerado com sucesso!")

if __name__ == "__main__":
    main()